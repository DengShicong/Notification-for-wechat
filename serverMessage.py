import os
import smtplib
import zipfile
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

import pandas as pd
import paramiko
import urllib.request
import json
import requests
from netmiko import ConnectHandler
from datetime import datetime
import threading
import time
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor

def send_wxpusher_message(app_token, content, summary, content_type, topic_ids, uids, url, verify_pay):
    wxpusher_url = "http://wxpusher.zjiecode.com/api/send/message"
    headers = {'Content-Type': 'application/json'}
    data = {
        "appToken": app_token,
        "content": content,
        "summary": summary,
        "contentType": content_type,
        "topicIds": topic_ids,
        "uids": uids,
        "url": url,
        "verifyPay": verify_pay
    }

    response = requests.post(wxpusher_url, headers=headers, data=json.dumps(data))
    return response.json()
def analyze_temperature_output(output):
    lines = output.split('\n')
    temperature_colors = {}  # 用于存储行和颜色的对应关系

    for line in lines:
        if '1/CMMA' in line:
            parts = line.split()  # 假设每个值都通过空格分隔
            current_temp = float(parts[1])  # 假设当前温度是第三个元素
            danger_temp = float(parts[-4])  # 假设危险温度是最后一个元素
            if current_temp < danger_temp:
                temperature_colors[line] = RGBColor(0, 128, 0)  # 绿色
            else:
                temperature_colors[line] = RGBColor(255, 0, 0)  # 红色

    return temperature_colors
def analyze_health_output(output):
    lines = output.split('\n')
    cpu_color = RGBColor(0, 0, 0)  # 默认黑色
    memory_color = {}  # 用于存储行和颜色的对应关系

    for line in lines:
        if 'CPU' in line:
            current_cpu_usage = float(line.split()[2])  # 假设当前使用率是第三个元素
            if current_cpu_usage < 80:
                cpu_color = RGBColor(0, 128, 0)  # 绿色
            else:
                cpu_color = RGBColor(255, 0, 0)  # 红色

        if 'Memory' in line:
            parts = line.split()  # 假设每个值都通过空格分隔
            current_memory_usage = float(parts[1])  # 假设当前使用率是第三个元素
            danger_memory_usage = float(parts[-2])  # 假设当前使用率是第三个元素
            if current_memory_usage < danger_memory_usage:
                memory_color = RGBColor(0, 128, 0)  # 绿色
            else:
                memory_color = RGBColor(255, 0, 0)  # 红色

    return cpu_color, memory_color
def sftp_down_file(host, user, password, server_path, local_path, timeout=10):
    try:
        t = paramiko.Transport((host, 22))
        t.banner_timeout = timeout
        t.connect(username=user, password=password)
        sftp = paramiko.SFTPClient.from_transport(t)
        sftp.get(server_path, local_path)
        t.close()
        return True
    except Exception as e:
        print(f"SFTP download failed for {host}: {e}")
        return False

def handle_device(row, log_dir, commands,app_token,uids):
    device_info = {
        "device_type": row["device_type"],
        "host": row["IP"],
        "username": row["username"],
        "password": row["password"],
        "port": int(row["port"]) if not pd.isna(row["port"]) else 22,
        "secret": row["secret"]
    }

    hostname = row.get("hostname")
    try:
        with ConnectHandler(**device_info) as net_connect:
            net_connect.enable()
            all_output = ""


            doc = Document()  # 创建 Word 文档对象
            doc.add_heading('健康度检测', level=1)  # 添加总标题
            doc.add_paragraph(f'设备信息: {row["IP"]}')
            doc.add_paragraph(f'时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'设备类型: {row["device_type"]}')
            doc.add_paragraph('巡检人员：')

            red_lines = []  # 存储带有红色字体的行
            is_normal = True  # 标记是否所有检查都正常

            for cmd in commands:
                output = net_connect.send_command(cmd)
                all_output += f"--- Command: {cmd} ---\n{output}\n\n"
                if output.strip():
                    doc.add_heading(f'Command: {cmd}', level=2)  # 将命令作为二级标题

                    if cmd == 'show health':
                        cpu_color, memory_color = analyze_health_output(output)

                    if cmd == 'show temperature':
                        temp_colors = analyze_temperature_output(output)

                    for line in output.split('\n'):
                        p = doc.add_paragraph(line)
                        p.style.font.name = 'Consolas'
                        p.style.font.size = Pt(10)

                        if cmd == 'show health':
                            if 'CPU' in line:
                                p.runs[0].font.color.rgb = cpu_color
                                if cpu_color == RGBColor(255, 0, 0):
                                    is_normal = False
                                    red_lines.append(line)
                            if 'Memory' in line:
                                p.runs[0].font.color.rgb = memory_color
                                if memory_color == RGBColor(255, 0, 0):
                                    is_normal = False
                                    red_lines.append(line)

                        if cmd == 'show temperature':
                            if line in temp_colors:
                                p.runs[0].font.color.rgb = temp_colors[line]
                                if temp_colors[line] == RGBColor(255, 0, 0):
                                    is_normal = False
                                    red_lines.append(line)


            content = f"巡检报告 - {device_info['host']}\n\n{all_output}"
            summary = f"巡检报告 - {device_info['host']}({hostname})"
            content_type = 3  # Markdown 格式
            response = send_wxpusher_message(app_token, content, summary, content_type, [], uids, "", False)
            print(f"WxPusher response for {device_info['host']}: {response}")


            doc.add_heading('巡检结果', level=1)
            doc.add_heading('设备状态总结：', level=2)
            doc.add_paragraph('正常' if is_normal else '警告')
            if not is_normal:
                doc.add_heading('存在的问题及建议：', level=2)
                for line in red_lines:
                    doc.add_paragraph(line)
            doc.add_heading('签名：', level=2)
            doc.add_paragraph('巡检人员签名：')

            if len(doc.paragraphs) > 1:
                timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
                filename = f"{timestamp}_{device_info['host']}.docx"
                filepath = os.path.join(log_dir, filename)
                doc.save(filepath)  # 保存 Word 文档

            sftp_local_path = os.path.join(log_dir, f"{timestamp}_{device_info['host']}_vcboot.cfg")
            sftp_remote_path = '/flash/working/vcboot.cfg'
            sftp_down_file(device_info['host'], device_info['username'], device_info['password'], sftp_remote_path,
                           sftp_local_path)
    except Exception as e:
        print(f"Connection failed for {device_info['host']}: {e}")

def execute_commands_and_save_logs(template_path, app_token, uids):
    template_data = pd.ExcelFile(template_path)
    assets_data = template_data.parse('assets')

    log_dir = "LOG-MESSAGE"
    os.makedirs(log_dir, exist_ok=True)

    threads = []
    for index, row in assets_data.iterrows():
        commands = template_data.parse(row["device_type"]).iloc[:, 1].dropna().tolist()
        thread = threading.Thread(target=handle_device, args=(row, log_dir, commands, app_token, uids))
        threads.append(thread)
        thread.start()

    for thread in threads:
        thread.join()

def compress_zip(source_dir,target_file):

    zipf = zipfile.ZipFile(target_file, 'w', zipfile.ZIP_DEFLATED)
    for root, dirs, files in os.walk(source_dir):
        for file in files:
            zipf.write(os.path.join(root, file))
    zipf.close()

source_dir = 'LOG-MESSAGE'
target_file = 'LOG-MESSAGE.zip'
compress_zip(source_dir,target_file)
def send():
    # 配置邮件信息
    time = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    sender = '2093967710@qq.com'
    receiver = ['dengshicong24@outlook.com','alanlan@139.com']
    subject = time+'巡检报告'
    attachment_path = 'LOG-MESSAGE.zip'
        # 构造邮件对象
    message = MIMEMultipart()
    message['From'] = sender
    message['To'] = ",".join(receiver)
    message['Subject'] = subject

        # 添加正文
    body = MIMEText('巡检报告')
    message.attach(body)

        # 添加附件
    with open(attachment_path, 'rb') as attachment:
        attachment_part = MIMEApplication(attachment.read())
        attachment_part.add_header('Content-Disposition', 'attachment', filename='LOG-MESSAGE.zip')
        message.attach(attachment_part)

        # 发送邮件

    with smtplib.SMTP_SSL('smtp.qq.com', 465) as server:
        server.ehlo()
        server.login('2093967710@qq.com', 'kfgoqjryeszkchcg')
        server.sendmail(sender, receiver, message.as_string())

        server.quit()
        print("邮件发送成功")

def main():
    template_path = 'template.xlsx'  # 替换为您的模板文件路径
    app_token = "AT_w39ZFDIeWh8oCbJeuayZBYSoxPKRcbhF"
    uids = ["UID_aHrUAEj5O8lSfSpisHRUllSWL0hx",
            "UID_4W43ksZPq65FCeC2E8tZsnGfn0sN",
            ]


    while True:
        print(f"Executing script at {datetime.now()}")
        execute_commands_and_save_logs(template_path, app_token, uids)
        send()
        print("Execution completed. Waiting for the next run...")
        time.sleep(600)  # 每10分钟运行一次


if __name__ == '__main__':
    main()